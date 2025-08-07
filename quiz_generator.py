import tkinter as tk
from tkinter import messagebox
from docx import Document
from docx.shared import Pt
import os
import subprocess

# 保存路径设为桌面
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
output_path = os.path.join(desktop, "output.docx")

questions = []

def add_question():
    q_text = question_entry.get("1.0", tk.END).strip()
    opts_text = options_entry.get("1.0", tk.END).strip()
    options = opts_text.split("\n")

    if not q_text or not options:
        messagebox.showwarning("Error", "Please enter both the question and options.")
        return

    questions.append({
        "question": q_text,
        "options": options
    })

    question_entry.delete("1.0", tk.END)
    options_entry.delete("1.0", tk.END)

    messagebox.showinfo("Success", "Question added successfully!")

def generate_docx():
    if not questions:
        messagebox.showwarning("No Questions", "Please add at least one question.")
        return

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'FangSong'  
    style.font.size = Pt(14)

    for idx, q in enumerate(questions, start=1):
        # 添加题干
        para_q = doc.add_paragraph()
        run_q = para_q.add_run(f"{idx}. {q['question']}")
        run_q.font.size = Pt(14)
        para_q.paragraph_format.space_after = Pt(6)
        para_q.paragraph_format.line_spacing = 1.5

        # 添加选项
        for opt in q['options']:
            para_opt = doc.add_paragraph(f"{opt}")
            para_opt.paragraph_format.left_indent = Pt(24)
            para_opt.paragraph_format.line_spacing = 1.5
            para_opt.paragraph_format.space_after = Pt(3)

        doc.add_paragraph("")  # 每题之间留空行

    doc.save(output_path)
    messagebox.showinfo("Completed", f"Word file saved to: {output_path}")

    # macOS 自动打开 Word 文件
    try:
        subprocess.call(["open", output_path])
    except:
        pass

# GUI 设置
root = tk.Tk()
root.title("Quiz Generator")
root.geometry("500x600")

# 输入题干
tk.Label(root, text="Question:").pack()
question_entry = tk.Text(root, height=4, width=60)
question_entry.pack()

# 输入选项
tk.Label(root, text="Options (one per line, e.g., A. xxx):").pack()
options_entry = tk.Text(root, height=6, width=60)
options_entry.pack()

# 功能按钮
tk.Button(root, text="Add Question", command=add_question).pack(pady=10)
tk.Button(root, text="Generate Word File", command=generate_docx).pack(pady=10)

# 底部标语
footer_label = tk.Label(
    root,
    text="work for future wealth",
    font=("Segoe Script", 12, "italic"),
    fg="gray"
)
footer_label.pack(side="bottom", pady=5)

# 启动界面主循环
root.mainloop()
